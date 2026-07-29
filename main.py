import sys

# Forzar codificación UTF-8 en consola de Windows para soporte completo de emojis y caracteres especiales
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8")
import os
import io
import re
import uuid
import unicodedata
import urllib.request
from difflib import SequenceMatcher
from pathlib import Path

from fastapi import FastAPI, HTTPException
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from groq import Groq
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

# ── GOOGLE DOC PRINCIPAL ─────────────────────────────────────────────
GOOGLE_DOC_ID = "1YQ1Z1f5YXhkZtQb8732QDD7wcU9LkeNLjuOA52WxdoU"
GOOGLE_DOC_URL = f"https://docs.google.com/document/d/{GOOGLE_DOC_ID}/export?format=txt"

# ── CLAVES DE API (GROQ & GEMINI) ───────────────────────────────────
# Las claves se configuran como variables de entorno en la plataforma de despliegue
_env_keys = [k.strip() for k in os.environ.get("GROQ_API_KEYS", "").split(",") if k.strip()]
GROQ_API_KEYS = _env_keys if _env_keys else []

GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "")

clientes_groq = [Groq(api_key=k) for k in GROQ_API_KEYS if k and "PEGA_AQUI" not in k]

try:
    import google.generativeai as genai
    genai.configure(api_key=GEMINI_API_KEY)
    cliente_gemini = genai.GenerativeModel("gemini-2.5-flash")
except Exception as e:
    print(f"⚠️ Error al inicializar cliente de Gemini: {e}")
    cliente_gemini = None

app = FastAPI(title="AI Document Assistant API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Servir archivos estáticos si existe la carpeta 'static'
static_path = Path("static")
if static_path.exists():
    app.mount("/static", StaticFiles(directory="static"), name="static")

# ── CARPETA CON DOCUMENTOS LOCALES OPCIONALES ───────────────────────
DOCS_FOLDER = Path("documentos")
DOCS_FOLDER.mkdir(exist_ok=True)

EXTENSIONES_SOPORTADAS = (".pdf", ".docx", ".xlsx", ".xls", ".csv")

# Almacén en memoria: { archivo_id: { tipo, origen, nombre, texto / hojas } }
archivos = {}


@app.get("/")
def index():
    if (Path("static") / "index.html").exists():
        return FileResponse("static/index.html")
    return {
        "status": "online",
        "service": "AI Assistant Backend",
        "documentos_cargados": len(archivos),
        "google_doc_conectado": True
    }


# ── Extracción de Google Doc ─────────────────────────────────────────
def descargar_google_doc(url: str = GOOGLE_DOC_URL) -> str:
    """Descarga el contenido de un Google Doc público exportado en formato TXT."""
    try:
        req = urllib.request.Request(
            url,
            headers={"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"}
        )
        with urllib.request.urlopen(req, timeout=15) as resp:
            contenido_bytes = resp.read()
            texto = contenido_bytes.decode("utf-8", errors="ignore")
            # Quitar BOM si existe
            if texto.startswith("\ufeff"):
                texto = texto[1:]
            return texto.strip()
    except Exception as e:
        print(f"⚠️  Error al descargar el Google Doc desde '{url}': {e}")
        return ""


# ── Extracción de texto de PDF ───────────────────────────────────────
def extraer_texto_pdf(contenido_bytes: bytes) -> str:
    if PdfReader is None:
        print("⚠️ 'pypdf' no está instalado. Omitiendo archivo PDF.")
        return ""
    lector = PdfReader(io.BytesIO(contenido_bytes))
    partes = []
    for i, pagina in enumerate(lector.pages):
        texto_pagina = pagina.extract_text() or ""
        if texto_pagina.strip():
            partes.append(f"[Página {i + 1}]\n{texto_pagina.strip()}")
    return "\n\n".join(partes)


# ── Extracción de Word ───────────────────────────────────────────────
def extraer_texto_word(contenido_bytes: bytes) -> str:
    if Document is None:
        print("⚠️ 'python-docx' no está instalado. Omitiendo archivo Word.")
        return ""
    doc = Document(io.BytesIO(contenido_bytes))
    partes = []

    for parrafo in doc.paragraphs:
        if parrafo.text.strip():
            partes.append(parrafo.text)

    for tabla in doc.tables:
        for fila in tabla.rows:
            celdas = [celda.text.strip() for celda in fila.cells]
            if any(celdas):
                partes.append(" | ".join(celdas))

    return "\n".join(partes)


# ── Extracción de CSV ────────────────────────────────────────────────
def extraer_dataframe_csv(contenido_bytes: bytes) -> pd.DataFrame:
    codificaciones = ["utf-8-sig", "utf-8", "latin-1", "cp1252"]

    ultimo_error = None
    for codificacion in codificaciones:
        try:
            texto_decodificado = contenido_bytes.decode(codificacion)
            df = pd.read_csv(
                io.StringIO(texto_decodificado),
                sep=None,
                engine="python",
            )
            return df
        except Exception as e:
            ultimo_error = e
            continue

    raise ValueError(f"No se pudo leer el CSV: {ultimo_error}")


# ── Carga de un archivo individual desde disco ───────────────────────
def cargar_archivo_desde_disco(ruta: Path) -> dict:
    extension = ruta.suffix.lower()
    contenido = ruta.read_bytes()

    if extension == ".pdf":
        texto = extraer_texto_pdf(contenido)
        return {"tipo": "pdf", "origen": "local", "nombre": ruta.name, "texto": texto}

    elif extension == ".docx":
        texto = extraer_texto_word(contenido)
        return {"tipo": "word", "origen": "local", "nombre": ruta.name, "texto": texto}

    elif extension in (".xlsx", ".xls"):
        excel = pd.ExcelFile(io.BytesIO(contenido))
        hojas_data = {}
        for hoja in excel.sheet_names:
            hojas_data[hoja] = pd.read_excel(excel, sheet_name=hoja)
        return {
            "tipo": "excel", "origen": "local", "nombre": ruta.name,
            "hojas": hojas_data, "hojas_nombres": excel.sheet_names,
        }

    elif extension == ".csv":
        df = extraer_dataframe_csv(contenido)
        return {
            "tipo": "csv", "origen": "local", "nombre": ruta.name,
            "hojas": {"CSV": df}, "hojas_nombres": ["CSV"],
            "filas": df.shape[0], "columnas": df.shape[1],
        }

    else:
        raise ValueError(f"Extensión no soportada: {extension}")


def construir_texto_excel(datos: dict) -> str:
    texto_completo = ""
    for nombre_hoja, df in datos["hojas"].items():
        df_str = df.copy()
        for col in df_str.columns:
            if df_str[col].dtype == object:
                df_str[col] = df_str[col].astype(str).str[:120]
        texto_completo += f"\n\n=== HOJA: {nombre_hoja} ===\n"
        texto_completo += df_str.to_string(index=False)
    return texto_completo


def obtener_texto_de_archivo(datos: dict) -> str:
    if datos["tipo"] in ("excel", "csv"):
        return construir_texto_excel(datos)
    return datos.get("texto", "")


# ══════════════════════════════════════════════════════════════════
# ── MOTOR RAG (Retrieval-Augmented Generation) ──────────────────────
# ══════════════════════════════════════════════════════════════════

TAMANO_CHUNK = 1200          # caracteres por fragmento — más grande para capturar más contexto
SOLAPAMIENTO_CHUNK = 400     # solapamiento amplio para no perder info entre fragmentos
TOP_K_CHUNKS = 18            # fragmentos recuperados por pregunta
TOP_K_RESUMEN = 30           # fragmentos para preguntas de tipo resumen/general
UMBRAL_RELEVANCIA = 0.01     # umbral muy bajo para no descartar nombres propios raros

VECTOR_STORE = {
    "chunks": [],       # [{ "texto": str, "documento": str }, ...]
    "vectorizer": None,  # TfidfVectorizer ya entrenado
    "matriz": None,      # matriz TF-IDF de todos los chunks
}


def dividir_en_chunks_rag(texto: str, tamano_chunk: int = TAMANO_CHUNK, solapamiento: int = SOLAPAMIENTO_CHUNK) -> list:
    texto = texto.strip()
    if not texto:
        return []

    chunks = []
    inicio = 0
    longitud = len(texto)

    while inicio < longitud:
        fin = min(inicio + tamano_chunk, longitud)

        if fin < longitud:
            corte = texto.rfind("\n", inicio, fin)
            if corte == -1 or corte <= inicio + int(tamano_chunk * 0.5):
                corte_espacio = texto.rfind(" ", inicio, fin)
                if corte_espacio > inicio:
                    corte = corte_espacio
            if corte != -1 and corte > inicio:
                fin = corte

        fragmento = texto[inicio:fin].strip()
        if fragmento:
            chunks.append(fragmento)

        siguiente_inicio = fin - solapamiento
        inicio = siguiente_inicio if siguiente_inicio > inicio else fin

    return chunks


def construir_indice_rag():
    global VECTOR_STORE

    todos_los_chunks = []
    for datos in archivos.values():
        texto_doc = obtener_texto_de_archivo(datos)
        for trozo in dividir_en_chunks_rag(texto_doc):
            todos_los_chunks.append({"texto": trozo, "documento": datos["nombre"]})

    if not todos_los_chunks:
        VECTOR_STORE = {"chunks": [], "vectorizer": None, "matriz": None}
        print("⚠️  No hay contenido para indexar en el motor RAG.")
        return

    textos = [c["texto"] for c in todos_los_chunks]
    vectorizer = TfidfVectorizer(
        max_df=0.9,
        min_df=1,
        ngram_range=(1, 2),
        sublinear_tf=True,
    )
    matriz = vectorizer.fit_transform(textos)

    VECTOR_STORE = {"chunks": todos_los_chunks, "vectorizer": vectorizer, "matriz": matriz}
    print(f"🔎 Índice RAG construido: {len(todos_los_chunks)} fragmento(s) de {len(archivos)} documento(s).")


PALABRAS_PUNTUALES = [
    "beneficiario", "beneficiaries", "beneficiarios", "partner", "socio", "headquarters",
    "head quarters", "sede", "pais", "paises", "country", "countries", "status", "estado",
    "categoria", "category", "presupuesto", "budget", "costo", "cost", "quien", "quienes",
    "cuanto", "cuantos", "how many", "who", "where", "donde", "cuando", "when", "cuales son sus",
    "cual es su", "what is the", "what are the", "cuales son los beneficiarios", "quienes son los beneficiarios"
]

PALABRAS_RESUMEN = [
    "resumen general", "resumen completo", "resume todo", "overview", "summary of all",
    "general summary", "give me a summary of all", "que sabes en general", "hablame de todo",
    "lista todos los proyectos", "list all projects", "summarize everything", "resumen de todo"
]


# ── PALABRAS NO ENTIDADES / ATRIBUTOS GENÉRICOS ──────────────────────
NON_ENTITY_WORDS = {
    # Español
    "beneficiario", "beneficiarios", "categoria", "categorias", "socio", "socios",
    "partner", "partners", "sede", "paises", "pais", "estado", "presupuesto", "costo",
    "costos", "quien", "quienes", "cuanto", "cuantos", "donde", "cuando", "cuales",
    "cual", "que", "descripcion", "solucion", "proyecto", "proyectos", "similar",
    "similares", "documento", "documentos", "informacion", "datos", "detalles", "resumen",
    "principales", "principal", "finales", "final", "objetivo", "objetivos",
    # Inglés
    "beneficiaries", "beneficiary", "category", "categories", "partner", "partners",
    "headquarters", "country", "countries", "status", "budget", "cost", "costs",
    "who", "where", "when", "which", "what", "description", "solution", "project",
    "projects", "similar", "document", "documents", "information", "info", "data",
    "details", "overview", "summary", "main", "final", "objective", "objectives"
}


def es_pregunta_general(texto: str) -> bool:
    """Detecta si la pregunta es de tipo resumen/general para recuperar más chunks."""
    t = texto.lower()
    if any(p in t for p in PALABRAS_PUNTUALES):
        return False
    return any(p in t for p in PALABRAS_RESUMEN)


def extraer_terminos_clave(pregunta: str) -> list:
    """Extrae términos de búsqueda clave de una pregunta — siglas, nombres propios, palabras significativas."""
    stopwords = {
        # Español
        "que", "cual", "cuales", "como", "donde", "cuando", "quien", "quienes", "cuanto",
        "cuantos", "por", "para", "sobre", "del", "de", "los", "las", "una", "uno", "unos", "unas",
        "con", "sin", "hay", "dame", "dime", "habla", "hablame", "decir", "dices", "sabes",
        "tienes", "tiene", "existe", "existen", "puede", "pueden", "es", "son", "fue", "eran",
        "esta", "estan", "sea", "sean", "ser", "trata", "tratan", "mas", "tambien", "todo", "toda",
        "todos", "todas", "informacion", "datos", "info", "algo", "nada", "ese", "esa", "esos",
        "esas", "este", "esta", "estos", "estas", "muy", "bien", "mal", "principales", "principal",
        "sus", "su", "el", "la", "los", "las", "un", "una",
        # Inglés
        "what", "whats", "where", "when", "who", "whom", "whose", "which", "why", "how",
        "is", "are", "was", "were", "be", "been", "being", "have", "has", "had", "do", "does",
        "did", "a", "an", "the", "and", "but", "if", "or", "because", "as", "until", "while",
        "of", "at", "by", "for", "with", "about", "against", "between", "into", "through",
        "during", "before", "after", "above", "below", "to", "from", "up", "down", "in",
        "out", "on", "off", "over", "under", "again", "further", "then", "once", "here",
        "there", "all", "any", "both", "each", "few", "more", "most", "other", "some",
        "such", "no", "nor", "not", "only", "own", "same", "so", "than", "too", "very",
        "s", "t", "can", "will", "just", "don", "should", "now", "tell", "give", "me",
        "show", "main", "final", "its", "their", "this", "that", "these", "those"
    }
    tokens = re.split(r'[\s,;.!?¿¡()\[\]"\'/\\]+', pregunta)
    terminos = []
    for t in tokens:
        t_limpio = t.strip()
        if len(t_limpio) < 2:
            continue
        t_lower = t_limpio.lower()
        if t_lower in stopwords:
            continue
        terminos.append(t_limpio)
    return terminos


def extraer_entidades_clave(pregunta: str) -> list:
    """Extrae únicamente nombres propios o entidades objetivo de la pregunta, descartando atributos genéricos."""
    terminos = extraer_terminos_clave(pregunta)
    return [t for t in terminos if t.lower() not in NON_ENTITY_WORDS]


def buscar_seccion_proyecto(pregunta: str) -> list:
    """
    Busca si la pregunta menciona un proyecto/entidad específico (ej. Casper, Tahaddi, SICAMEX)
    y extrae la sección COMPLETA dedicada a dicho proyecto directamente del documento.
    """
    entidades = extraer_entidades_clave(pregunta)
    if not entidades:
        return []

    resultados = []
    vistos = set()

    for datos in archivos.values():
        texto_doc = obtener_texto_de_archivo(datos)
        nombre_doc = datos["nombre"]

        for entidad in entidades:
            if len(entidad) < 3:
                continue

            # 1. Buscar primero un encabezado numerado explícito: ej. "20. Casper", "1. Kaikaia", "17.SIMN"
            patron_encabezado = re.compile(rf'(?i)(?:^|\n)\s*(\d+\.\s*{re.escape(entidad)}\b[\s\S]*?)(?=\n\s*\d+\.|\Z)')
            matches = list(patron_encabezado.finditer(texto_doc))

            if not matches:
                # 2. Si no hay encabezado numerado exacto, buscar título de línea sola o nombre de proyecto
                patron_general = re.compile(rf'(?i)(?:^|\n)\s*({re.escape(entidad)}\b[\s\S]*?)(?=\n\s*\d+\.|\n\s*Project|\Z)')
                matches = list(patron_general.finditer(texto_doc))

            for match in matches:
                bloque = match.group(1).strip()
                bloque_lower = bloque.lower()
                if len(bloque) > 80 and any(k in bloque_lower for k in ("category", "partner", "beneficiar", "head quarters", "project status", "description")):
                    bloque_util = bloque[:3500]
                    clave = bloque_util[:100]
                    if clave not in vistos:
                        vistos.add(clave)
                        resultados.append({
                            "texto": bloque_util,
                            "documento": nombre_doc,
                            "score": 500.0,  # Máxima prioridad para la sección oficial del proyecto
                            "fuente": "seccion_proyecto",
                        })

    return resultados


def buscar_en_texto_completo(pregunta: str, ventana_chars: int = 1500) -> list:
    """
    Búsqueda de último recurso: escanea el texto COMPLETO de cada documento
    buscando los términos clave de la pregunta y devuelve ventanas de contexto.
    """
    terminos = extraer_entidades_clave(pregunta) or extraer_terminos_clave(pregunta)
    if not terminos:
        return []

    resultados = []
    vistos = set()

    for datos in archivos.values():
        texto_doc = obtener_texto_de_archivo(datos)
        nombre_doc = datos["nombre"]
        texto_lower = texto_doc.lower()

        for termino in terminos:
            t_lower = termino.lower()
            pos = 0
            while True:
                idx = texto_lower.find(t_lower, pos)
                if idx == -1:
                    break
                inicio = max(0, idx - ventana_chars // 2)
                fin = min(len(texto_doc), idx + ventana_chars // 2)
                fragmento = texto_doc[inicio:fin].strip()
                clave = fragmento[:100]
                if clave not in vistos:
                    vistos.add(clave)
                    resultados.append({
                        "texto": fragmento,
                        "documento": nombre_doc,
                        "score": 2.0 if t_lower in fragmento.lower() else 0.5,
                        "fuente": "busqueda_completa",
                    })
                pos = idx + 150  # avanzar para no generar ventanas idénticas

    return resultados


def buscar_por_keyword(pregunta: str, top_k: int = TOP_K_CHUNKS) -> list:
    """Búsqueda exacta por palabras clave filtrando stopwords y atributos genéricos."""
    chunks = VECTOR_STORE.get("chunks", [])
    if not chunks:
        return []

    palabras = [p.lower() for p in (extraer_entidades_clave(pregunta) or extraer_terminos_clave(pregunta))]
    if not palabras:
        return []

    resultados = []
    for chunk in chunks:
        texto_lower = chunk["texto"].lower()
        coincidencias = sum(1 for p in palabras if p in texto_lower)
        if coincidencias > 0:
            score_kw = coincidencias / len(palabras)
            # Boost extra si coincide una entidad específica
            for p in palabras:
                if len(p) >= 4 and p in texto_lower:
                    score_kw += 1.0
            resultados.append({"texto": chunk["texto"], "documento": chunk["documento"], "score": score_kw, "kw": True})

    resultados.sort(key=lambda x: x["score"], reverse=True)
    return resultados[:top_k]


def buscar_chunks_relevantes(pregunta: str, top_k: int = TOP_K_CHUNKS, umbral_minimo: float = UMBRAL_RELEVANCIA) -> list:
    """
    Búsqueda híbrida en 4 capas:
      0. Extracción directa de sección oficial del proyecto (máxima prioridad)
      1. TF-IDF semántico sobre chunks indexados con boost por entidad
      2. Coincidencia exacta de palabras clave en chunks
      3. Búsqueda literal completa en el texto íntegro del documento
    Garantiza que el proyecto buscado siempre aparezca primero en los fragmentos.
    """
    chunks = VECTOR_STORE.get("chunks", [])
    es_general = es_pregunta_general(pregunta)
    entidades_clave = extraer_entidades_clave(pregunta)
    terminos_clave = extraer_terminos_clave(pregunta)

    vistos = set()
    resultados = []

    # ── CAPA 0: Extracción directa de sección del proyecto ─────────────
    secciones_proyecto = buscar_seccion_proyecto(pregunta)
    for r in secciones_proyecto:
        clave = r["texto"][:100]
        if clave not in vistos:
            vistos.add(clave)
            resultados.append(r)

    # ── CAPA 1: TF-IDF semántico ─────────────────────────────────────
    if chunks and VECTOR_STORE.get("vectorizer"):
        vector_pregunta = VECTOR_STORE["vectorizer"].transform([pregunta])
        similitudes = cosine_similarity(vector_pregunta, VECTOR_STORE["matriz"])[0]
        indices_ordenados = similitudes.argsort()[::-1]

        for idx in indices_ordenados[:top_k * 3]:
            score = float(similitudes[idx])
            chunk = chunks[idx]
            texto_chunk_lower = chunk["texto"].lower()

            # Boost por coincidencia de entidad explícita
            for ent in entidades_clave:
                if len(ent) >= 3 and ent.lower() in texto_chunk_lower:
                    score += 50.0
                    if re.search(rf'(?i)\b\d+\.\s*{re.escape(ent)}\b', chunk["texto"]):
                        score += 100.0

            # Boost si coincide con términos clave de la pregunta
            for term in terminos_clave:
                if len(term) >= 4 and term.lower() in texto_chunk_lower:
                    score += 2.0

            if score < umbral_minimo and not resultados:
                break
            clave = chunk["texto"][:100]
            if clave not in vistos:
                vistos.add(clave)
                resultados.append({"texto": chunk["texto"], "documento": chunk["documento"], "score": score})

    # ── CAPA 2: Keyword exacto ──────────────────────────────────────
    kw_resultados = buscar_por_keyword(pregunta, top_k=top_k * 2)
    for r in kw_resultados:
        clave = r["texto"][:100]
        if clave not in vistos:
            vistos.add(clave)
            resultados.append(r)

    # ── CAPA 3: Búsqueda literal completa ────────────────────────────
    texto_completo_resultados = buscar_en_texto_completo(pregunta)
    for r in texto_completo_resultados:
        clave = r["texto"][:100]
        if clave not in vistos:
            vistos.add(clave)
            resultados.append(r)

    if not resultados:
        return []

    # Si la pregunta nombra entidades específicas, priorizar fuertemente los fragmentos que las contengan
    if entidades_clave:
        for ent in entidades_clave:
            ent_lower = ent.lower()
            for r in resultados:
                if ent_lower in r["texto"].lower():
                    r["score"] += 200.0

    # Ordenar por score descendente
    resultados.sort(key=lambda x: x["score"], reverse=True)

    limite = top_k if not es_general else max(top_k, TOP_K_RESUMEN)
    return resultados[:limite]


# ── Descarga Google Doc y escanea carpeta DOCS_FOLDER ────────────────
def cargar_documentos_desde_carpeta():
    archivos.clear()

    # 1. Cargar Google Doc
    print("🌐 Descargando contenido actualizado del Google Doc...")
    texto_gdoc = descargar_google_doc()
    if texto_gdoc:
        id_gdoc = "gdoc_principal"
        archivos[id_gdoc] = {
            "tipo": "google_doc",
            "origen": "google_docs",
            "nombre": "Documento Principal (Google Doc)",
            "texto": texto_gdoc,
        }
        print(f"✅ Google Doc cargado exitosamente ({len(texto_gdoc)} caracteres).")
    else:
        print("⚠️  No se pudo descargar el Google Doc. Se usarán solo archivos locales si existen.")

    # 2. Cargar archivos locales adicionales de DOCS_FOLDER
    rutas = sorted(
        p for p in DOCS_FOLDER.iterdir()
        if p.is_file() and p.suffix.lower() in EXTENSIONES_SOPORTADAS
    )

    for ruta in rutas:
        try:
            datos = cargar_archivo_desde_disco(ruta)
            archivo_id = str(uuid.uuid4())[:8]
            archivos[archivo_id] = datos
        except Exception as e:
            print(f"⚠️  No se pudo cargar '{ruta.name}': {e}")

    print(f"✅ Total: {len(archivos)} documento(s) cargado(s).")

    # Reconstruir índice RAG
    construir_indice_rag()

    # Limpiar memoria cache para forzar respuestas actualizadas
    MEMORIA_QA.clear()


@app.on_event("startup")
def al_iniciar():
    cargar_documentos_desde_carpeta()


# ── Recargar documentos manualmente ──────────────────────────────────
@app.post("/recargar-documentos")
def recargar_documentos():
    cargar_documentos_desde_carpeta()
    return {"ok": True, "total": len(archivos), "chunks_indexados": len(VECTOR_STORE["chunks"])}


# ── Listar documentos cargados ───────────────────────────────────────
@app.get("/archivos")
def listar_archivos():
    lista = []
    for archivo_id, datos in archivos.items():
        item = {
            "archivo_id": archivo_id,
            "nombre": datos["nombre"],
            "tipo": datos["tipo"],
            "origen": datos.get("origen", "local"),
        }
        if datos["tipo"] == "excel":
            item["hojas"] = datos["hojas_nombres"]
        elif datos["tipo"] == "csv":
            item["filas"] = datos.get("filas")
            item["columnas"] = datos.get("columnas")
        lista.append(item)
    return {"ok": True, "archivos": lista, "total": len(lista), "chunks_indexados": len(VECTOR_STORE["chunks"])}


# ── MEMORIA DE PREGUNTAS Y RESPUESTAS ───────────────────────────────
CLAVE_MEMORIA_GLOBAL = "TODOS_LOS_DOCUMENTOS"
MEMORIA_QA = {}
MEMORIA_MAX_POR_CLAVE = 60

UMBRAL_SIMILITUD_TEXTO = 0.88
UMBRAL_SIMILITUD_PALABRAS = 0.7
UMBRAL_DIFERENCIA_LONGITUD = 0.4


def normalizar_texto(texto: str) -> str:
    texto = texto.strip().lower()
    texto = unicodedata.normalize("NFKD", texto)
    texto = "".join(c for c in texto if not unicodedata.combining(c))
    texto = re.sub(r"[^a-z0-9ñ\s]", " ", texto)
    texto = re.sub(r"\s+", " ", texto).strip()
    return texto


def similitud_preguntas(a_normalizada: str, b_normalizada: str) -> tuple:
    similitud_texto = SequenceMatcher(None, a_normalizada, b_normalizada).ratio()

    palabras_a = set(a_normalizada.split())
    palabras_b = set(b_normalizada.split())
    if not palabras_a or not palabras_b:
        similitud_palabras = 0.0
    else:
        interseccion = palabras_a & palabras_b
        union = palabras_a | palabras_b
        similitud_palabras = len(interseccion) / len(union)

    return similitud_texto, similitud_palabras


def buscar_en_memoria(clave: str, pregunta_texto: str):
    entradas = MEMORIA_QA.get(clave, [])
    if not entradas:
        return None

    pregunta_normalizada = normalizar_texto(pregunta_texto)
    longitud_actual = len(pregunta_normalizada)
    mejor_respuesta = None
    mejor_puntaje = 0.0

    for entrada in entradas:
        longitud_guardada = len(entrada["pregunta_normalizada"])
        if longitud_actual == 0 or longitud_guardada == 0:
            continue

        diferencia_longitud = abs(longitud_actual - longitud_guardada) / max(longitud_actual, longitud_guardada)
        if diferencia_longitud > UMBRAL_DIFERENCIA_LONGITUD:
            continue

        sim_texto, sim_palabras = similitud_preguntas(pregunta_normalizada, entrada["pregunta_normalizada"])

        if sim_texto >= UMBRAL_SIMILITUD_TEXTO and sim_palabras >= UMBRAL_SIMILITUD_PALABRAS:
            if sim_texto > mejor_puntaje:
                mejor_puntaje = sim_texto
                mejor_respuesta = entrada["respuesta"]

    return mejor_respuesta


def guardar_en_memoria(clave: str, pregunta_texto: str, respuesta: str):
    entradas = MEMORIA_QA.setdefault(clave, [])

    entradas.append({
        "pregunta": pregunta_texto,
        "pregunta_normalizada": normalizar_texto(pregunta_texto),
        "respuesta": respuesta,
    })

    if len(entradas) > MEMORIA_MAX_POR_CLAVE:
        MEMORIA_QA[clave] = entradas[-MEMORIA_MAX_POR_CLAVE:]


def consultar_groq(prompt: str, temperatura: float = 0.15) -> str:
    # 1. Intentar con la clave de Groq
    if clientes_groq:
        for indice, cliente in enumerate(clientes_groq):
            try:
                respuesta = cliente.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=[{"role": "user", "content": prompt}],
                    max_tokens=2048,
                    temperature=temperatura,
                )
                return respuesta.choices[0].message.content
            except Exception as e:
                print(f"⚠️ Aviso: Error con Groq (clave {indice + 1}): {e}")

    # 2. Respaldo inteligente con la clave de Gemini
    if cliente_gemini:
        try:
            respuesta_gemini = cliente_gemini.generate_content(prompt)
            if respuesta_gemini and respuesta_gemini.text:
                return respuesta_gemini.text
        except Exception as e:
            print(f"⚠️ Aviso: Error con Gemini: {e}")
            raise e

    raise Exception("No se pudo obtener respuesta del modelo de IA.")


class Pregunta(BaseModel):
    texto: str


PALABRAS_CLAVE_FUENTE = [
    "en que documento", "en que archivo", "en que pdf", "en que parte",
    "en que pagina", "en que hoja", "de que documento", "de donde sacaste",
    "de donde salio", "de donde viene", "cual es la fuente", "cual es tu fuente",
    "que documento", "que archivo", "donde dice eso", "donde sale eso",
    "cita la fuente", "cita el documento", "menciona el documento",
    "menciona la fuente", "indica la fuente", "indica el documento",
]


def pregunta_pide_fuente(texto_pregunta: str) -> bool:
    texto_normalizado = normalizar_texto(texto_pregunta)
    return any(clave in texto_normalizado for clave in PALABRAS_CLAVE_FUENTE)


_PATRON_CITA_CLAUSULA = re.compile(
    r"[,;]?\s*(según|como se (menciona|indica|señala) en|de acuerdo (a|con)|"
    r"tal como (se )?(menciona|indica) en)\s+[\"'“]?[\w áéíóúñÁÉÍÓÚÑ\-\.]{0,80}?"
    r"\.(pdf|docx|xlsx|xls|csv)[\"'”]?(\s*\[?\(?[Pp]ágina\s*\d+\)?\]?)?",
    re.IGNORECASE,
)
_PATRON_ARCHIVO_SUELTO = re.compile(
    r"[\"'“]?\b[\w\-]+\.(pdf|docx|xlsx|xls|csv)[\"'”]?", re.IGNORECASE
)
_PATRON_PAGINA_SUELTA = re.compile(r"\[?\(?[Pp]ágina\s*\d+\)?\]?")

_PATRON_SCRIPT_EXTRANO = re.compile(
    r"[\u4e00-\u9fff\u3040-\u30ff\u30a0-\u30ff\uac00-\ud7af\u0400-\u04FF\u0600-\u06FF\u0e00-\u0e7f]+"
)


def limpiar_texto_general(texto: str) -> str:
    limpio = _PATRON_SCRIPT_EXTRANO.sub("", texto)
    limpio = re.sub(r"[ \t]{2,}", " ", limpio)
    limpio = re.sub(r"\n{3,}", "\n\n", limpio)
    limpio = re.sub(r"\s+([.,;:])", r"\1", limpio)
    return limpio.strip()


def limpiar_menciones_a_documentos(texto: str) -> str:
    limpio = _PATRON_CITA_CLAUSULA.sub("", texto)
    limpio = _PATRON_ARCHIVO_SUELTO.sub("", limpio)
    limpio = _PATRON_PAGINA_SUELTA.sub("", limpio)

    limpio = re.sub(r"\s+([.,;:])", r"\1", limpio)
    limpio = re.sub(r"[ \t]{2,}", " ", limpio)
    limpio = re.sub(r"\n{3,}", "\n\n", limpio)
    limpio = re.sub(r"\(\s*\)", "", limpio)
    limpio = re.sub(r"«\s*»", "", limpio)
    return limpio.strip()


def generar_respuesta_ia(texto_pregunta: str) -> str:
    chunks_relevantes = buscar_chunks_relevantes(texto_pregunta)

    if not chunks_relevantes:
        return "No information found related to your question in the loaded document. Try rephrasing or verify the document contains that information."

    citar_fuentes = pregunta_pide_fuente(texto_pregunta)

    if citar_fuentes:
        contexto = "\n\n".join(
            f'[Fragmento {i+1} · documento: "{c["documento"]}"]\n{c["texto"]}'
            for i, c in enumerate(chunks_relevantes)
        )
        instruccion_citas = (
            '- The user IS asking for the source: for each important piece of data, '
            'indicate the exact name of the document it comes from.'
        )
    else:
        contexto = "\n\n".join(
            f"[Fragmento {i+1}]\n{c['texto']}"
            for i, c in enumerate(chunks_relevantes)
        )
        instruccion_citas = (
            '- DO NOT mention file names or say things like "according to the document...", '
            '"as mentioned in...". Give the information as a direct answer, without citing sources.'
        )

    # Determinar si la pregunta es general/resumen para ajustar instrucciones
    es_general = es_pregunta_general(texto_pregunta)
    instruccion_alcance = (
        "- GENERAL OVERVIEW REQUEST: Use ALL provided fragments to build a complete, well-structured overview. "
        "Organize by topics or projects if multiple exist."
        if es_general else
        "- STRICT MANDATORY RULE FOR SPECIFIC/PUNTUAL QUESTIONS: "
        "Answer ONLY AND EXACTLY what was explicitly asked. Be direct, punctual, and concise. "
        "Do NOT add general project descriptions, solution details, categories, project status, headquarters, partner, similar projects, or unasked metadata. "
        "For example, if the user asks 'What are the final beneficiaries of Casper?', answer ONLY with the final beneficiaries of Casper (e.g. '+4,900 Children and Families') in 1 short sentence or bullet point. "
        "Do NOT output a full project profile, description, solution breakdown, or similar projects list unless explicitly asked."
    )

    prompt = f"""You are an expert document analysis assistant. Your job is to answer questions based strictly on the document fragments provided below.

{len(chunks_relevantes)} fragments were retrieved from the document using semantic search, keyword matching, and full-text literal search.

DOCUMENT FRAGMENTS ({len(chunks_relevantes)} retrieved):
{contexto}

USER QUESTION: {texto_pregunta}

MANDATORY RULES — FOLLOW THESE EXACTLY:

1. Read every fragment carefully before answering.

2. LANGUAGE REQUIREMENT: ALWAYS respond 100% IN ENGLISH. Even if the user's question is written in Spanish, French, or another language, your entire response MUST BE WRITTEN IN ENGLISH.

3. DIRECT & CONCISE ANSWERING:
{instruccion_alcance}

4. ACCURACY: Only use facts explicitly present in the fragments. Never invent data or add unrequested extra sections.

5. MISSING DATA: Only say the document does not contain information if, after reviewing ALL fragments, the topic does not appear at all.

{instruccion_citas}
6. Format cleanly and directly in English."""

    respuesta = consultar_groq(prompt)
    respuesta = limpiar_texto_general(respuesta)

    if not citar_fuentes:
        respuesta = limpiar_menciones_a_documentos(respuesta)

    return respuesta


@app.post("/preguntar")
async def preguntar(pregunta: Pregunta):
    if not archivos:
        raise HTTPException(status_code=400, detail="No hay documentos cargados en el servidor todavía.")

    try:
        respuesta_en_memoria = buscar_en_memoria(CLAVE_MEMORIA_GLOBAL, pregunta.texto)
        if respuesta_en_memoria is not None:
            return {
                "ok": True,
                "respuesta": respuesta_en_memoria,
                "desde_memoria": True,
            }

        respuesta_final = generar_respuesta_ia(pregunta.texto)
        guardar_en_memoria(CLAVE_MEMORIA_GLOBAL, pregunta.texto, respuesta_final)

        return {
            "ok": True,
            "respuesta": respuesta_final,
            "desde_memoria": False,
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error con Groq: {str(e)}")


@app.get("/depurar-rag")
def depurar_rag(pregunta: str, top_k: int = 5):
    resultados = buscar_chunks_relevantes(pregunta, top_k=top_k, umbral_minimo=0.0)
    return {
        "ok": True,
        "pregunta": pregunta,
        "total_chunks_indexados": len(VECTOR_STORE["chunks"]),
        "resultados": [
            {
                "documento": r["documento"],
                "score": round(r["score"], 4),
                "fragmento": r["texto"][:300],
            }
            for r in resultados
        ],
    }


@app.get("/memoria")
def ver_memoria():
    entradas = MEMORIA_QA.get(CLAVE_MEMORIA_GLOBAL, [])
    return {
        "ok": True,
        "total_guardadas": len(entradas),
        "preguntas": [e["pregunta"] for e in entradas],
    }


@app.delete("/memoria")
def limpiar_memoria():
    MEMORIA_QA.pop(CLAVE_MEMORIA_GLOBAL, None)
    return {"ok": True}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
